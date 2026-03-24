"""Co-Scientist Agent 보고서 생성 — 기술 보고서 + 경영진 보고서 (Word)."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
DOCS_DIR = os.path.join(BASE_DIR, "docs")
os.makedirs(DOCS_DIR, exist_ok=True)

DIAGRAM_PATH = os.path.join(DOCS_DIR, "architecture.png")


def set_cell_shading(cell, color_hex):
    """테이블 셀 배경색 설정."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shading_elm)


def style_heading(doc, text, level=1):
    """스타일이 적용된 Heading 추가."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_table(doc, headers, rows, col_widths=None):
    """스타일이 적용된 테이블 추가."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "1565C0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ═══════════════════════════════════════════════════════════
#  보고서 1: 기술 보고서 (전문가용)
# ═══════════════════════════════════════════════════════════
def generate_technical_report():
    doc = Document()

    # 제목
    title = doc.add_heading("Co-Scientist Agent\n기술 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Display R&D 논문 지원 시스템 — 아키텍처 및 기술 상세")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("작성일: 2026-03-24  |  버전: 0.2.0")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── 목차 ──
    style_heading(doc, "목차", 1)
    toc_items = [
        "1. 시스템 개요",
        "2. 아키텍처 다이어그램",
        "3. 기술 스택",
        "4. Supervisor 파이프라인",
        "5. 14개 에이전트 상세",
        "6. 데이터 파이프라인",
        "7. 날짜 파싱 엔진",
        "8. 평가 결과",
        "9. 향후 계획",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── 1. 시스템 개요 ──
    style_heading(doc, "1. 시스템 개요", 1)
    doc.add_paragraph(
        "Co-Scientist Agent는 Display R&D 분야의 논문 데이터를 MariaDB(관계형)와 "
        "Milvus(벡터 DB)에 적재하고, LangGraph 기반 14개 전문 에이전트가 사용자 질문에 "
        "답변하는 RAG(Retrieval-Augmented Generation) 시스템이다."
    )
    doc.add_paragraph(
        "핵심 설계 원칙:\n"
        "• OpenAI-compatible API 통일 — 폐쇄망 이전 시 .env만 변경\n"
        "• Hybrid 검색 — Dense(IVF_FLAT/IP) + Sparse(BM25) 앙상블\n"
        "• 한국어 날짜 파싱 — Regex 기반 D1~D4 패턴, 서버 시간 상대 계산\n"
        "• 멀티턴 대화 — 최대 5턴 히스토리, 이전 턴 논문 제목 자동 추적\n"
        "• 트레이싱 — Langfuse 기반 전체 파이프라인 관측"
    )

    # ── 2. 아키텍처 다이어그램 ──
    style_heading(doc, "2. 아키텍처 다이어그램", 1)
    if os.path.exists(DIAGRAM_PATH):
        doc.add_picture(DIAGRAM_PATH, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[다이어그램 파일 없음 — scripts/generate_diagram.py 실행 필요]")

    # ── 3. 기술 스택 ──
    style_heading(doc, "3. 기술 스택", 1)
    add_table(doc,
        ["구분", "기술", "비고"],
        [
            ["Backend", "FastAPI + Uvicorn", "Port 20035, async"],
            ["Agent Framework", "LangGraph (StateGraph)", "조건부 라우팅"],
            ["LLM", "vLLM / LM Studio", "OpenAI-compat API"],
            ["Embedding", "TEI / LM Studio", "OpenAI-compat API"],
            ["Vector DB", "Milvus 2.x", "IVF_FLAT/IP + BM25"],
            ["RDBMS", "MariaDB 10.x", "SQLAlchemy ORM"],
            ["Tracing", "Langfuse", "LLM 관측/디버깅"],
            ["UI", "Open WebUI", "OpenAI-compat 연동"],
            ["Python", "3.12 (WinPython)", "독립 환경"],
        ],
        col_widths=[3, 5, 6],
    )

    # ── 4. Supervisor 파이프라인 ──
    style_heading(doc, "4. Supervisor 파이프라인", 1)
    doc.add_paragraph(
        "모든 사용자 질문은 Supervisor 파이프라인(LangGraph StateGraph)을 거쳐 처리된다."
    )
    add_table(doc,
        ["단계", "함수", "처리 내용", "사용 기술"],
        [
            ["①", "extract_dates", "한국어 날짜 표현 → coverdate 필터 범위 변환", "Regex (D1~D4)"],
            ["②", "extract_conditions", "키워드/저자/DOI/volume/issue 추출", "LLM (JSON 출력)"],
            ["③", "classify_intent", "14개 에이전트 중 하나로 의도 분류", "LLM 분류기"],
            ["④", "route_to_agent", "분류 결과에 따라 해당 에이전트 실행", "LangGraph 조건부 엣지"],
            ["⑤", "add_citation", "출처 논문 목록 + 저작권 표시 추가", "후처리"],
        ],
        col_widths=[1.2, 3.5, 5, 4],
    )

    # ── 5. 14개 에이전트 ──
    style_heading(doc, "5. 14개 에이전트 상세", 1)

    style_heading(doc, "Phase 1: 기본 검색/분석", 2)
    add_table(doc,
        ["에이전트", "기능", "데이터 소스"],
        [
            ["Paper QA", "논문 내용 검색 및 Q&A", "Milvus + MariaDB"],
            ["Literature Survey", "주제별 문헌 리뷰 종합", "Milvus"],
            ["Paper Deep Dive", "특정 논문 심층 분석", "MariaDB"],
            ["Analytics", "통계/집계/트렌드 (연도별, 저자별 등)", "MariaDB SQL"],
        ],
        col_widths=[3.5, 6, 4],
    )

    style_heading(doc, "Phase 2: 아이디어 발굴", 2)
    add_table(doc,
        ["에이전트", "기능", "데이터 소스"],
        [
            ["Idea Generator", "논문 간 교차 분석으로 연구 아이디어 제안", "Milvus"],
            ["Cross Domain", "타 분야 기술의 디스플레이 적용 가능성 분석", "Milvus"],
            ["Trend Analyzer", "기술 카테고리별 트렌드 분석 (12개 분류)", "Milvus"],
        ],
        col_widths=[3.5, 7, 3],
    )

    style_heading(doc, "Phase 3: 실험/전략", 2)
    add_table(doc,
        ["에이전트", "기능", "데이터 소스"],
        [
            ["Experiment Planner", "문헌 기반 실험 설계 제안", "Milvus"],
            ["Material Advisor", "재료/공정 비교 분석", "Milvus"],
            ["Patent Landscaper", "특허 동향 분석 및 공백 식별", "Milvus"],
            ["Competitive Intel", "경쟁사 연구 동향 모니터링", "Milvus"],
        ],
        col_widths=[3.5, 6, 4],
    )

    style_heading(doc, "Phase 4: 산출물 생성", 2)
    add_table(doc,
        ["에이전트", "기능", "데이터 소스"],
        [
            ["Report Drafter", "연구 보고서/발표자료 초안 생성", "Milvus"],
            ["Peer Review", "가상 동료 리뷰 피드백 제공", "Milvus"],
            ["Knowledge Connector", "논문 저자 기반 전문가 매칭", "Milvus"],
        ],
        col_widths=[3.5, 6, 4],
    )

    # ── 6. 데이터 파이프라인 ──
    style_heading(doc, "6. 데이터 파이프라인", 1)
    doc.add_paragraph(
        "데이터 적재는 2단계로 수행된다:\n\n"
        "① CSV → MariaDB (load_csv_to_mariadb.py)\n"
        "   • sample_paper.csv (1,614편) → sid_v_09_01 테이블\n"
        "   • 17개 컬럼: doi, coverdate(INT64 YYYYMMDD), title, paper_text 등\n\n"
        "② MariaDB → Milvus (load_mariadb_to_milvus.py)\n"
        "   • 논문 텍스트를 chunk 단위로 분할 (chunker.py)\n"
        "   • 각 chunk에 Dense embedding (OpenAI-compat API) 생성\n"
        "   • Milvus에 Dense(IVF_FLAT/IP) + Sparse(BM25) 인덱스로 적재\n"
        "   • ⚠️ 논문 편수 집계 시 chunk_id=1만 카운트해야 함"
    )

    # ── 7. 날짜 파싱 ──
    style_heading(doc, "7. 날짜 파싱 엔진 (date_parser.py)", 1)
    add_table(doc,
        ["패턴", "예시", "변환 결과"],
        [
            ["D1 (절대-특정)", "2024년 11월", "20241101 ~ 20241130"],
            ["D2 (절대-범위)", "2023년 3분기 / 2022~2024년", "20230701~20230930 / 20220101~20241231"],
            ["D3 (상대)", "작년 여름 / 최근 6개월", "서버 시간 기준 자동 계산"],
            ["D4 (비교)", "2020년과 2024년 비교", "20200101 ~ 20241231"],
        ],
        col_widths=[3, 5, 6],
    )
    doc.add_paragraph(
        "• 2자리 연도 자동 변환: 25년→2025년, 99년→1999년\n"
        "• 테스트: 389건 정규 테스트 + 1,155건 확장 테스트 — 100% 통과"
    )

    # ── 8. 평가 결과 ──
    style_heading(doc, "8. 평가 결과 (2026-03-20 기준)", 1)
    doc.add_paragraph("기준 모델: qwen3-0.6b (경량 모델)")
    add_table(doc,
        ["평가 항목", "테스트 건수", "정확도", "비고"],
        [
            ["날짜 파싱", "1,155건", "100%", "Regex 기반, 모델 무관"],
            ["필터 적용", "85건", "100%", "날짜 → SQL WHERE 변환"],
            ["Intent 분류", "235건", "50%", "0.6B 모델 한계, 235B에서 개선 예상"],
            ["E2E 파이프라인", "10건", "70%", "전체 흐름 샘플 테스트"],
        ],
        col_widths=[3, 2.5, 2, 6],
    )

    # ── 9. 향후 계획 ──
    style_heading(doc, "9. 향후 계획", 1)
    doc.add_paragraph(
        "• 대형 모델(235B) 전환으로 Intent 분류 정확도 향상\n"
        "• 폐쇄망 배포 — .env 전환만으로 vLLM + TEI 환경 이전\n"
        "• QA 테스트셋 확장 (2,080건 → 자동 생성 확대)\n"
        "• 에이전트 간 협업 (멀티 에이전트 체이닝)\n"
        "• 실시간 논문 업데이트 파이프라인 구축"
    )

    path = os.path.join(DOCS_DIR, "기술_보고서_Co-Scientist_Agent.docx")
    doc.save(path)
    print(f"Technical report saved: {path}")


# ═══════════════════════════════════════════════════════════
#  보고서 2: 경영진 보고서 (비전문가용)
# ═══════════════════════════════════════════════════════════
def generate_executive_report():
    doc = Document()

    # 제목
    title = doc.add_heading("AI 논문 분석 비서\n경영진 보고서", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("디스플레이 R&D 연구원을 위한 AI 기반 논문 검색·분석 도우미")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("보고일: 2026-03-24")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ── 한 줄 요약 ──
    style_heading(doc, "한 줄 요약", 1)
    p = doc.add_paragraph()
    run = p.add_run(
        "연구원이 채팅창에 질문만 입력하면, AI가 1,600편 이상의 논문 데이터베이스에서 "
        "관련 논문을 찾아 분석하고, 한국어로 답변해주는 시스템입니다."
    )
    run.font.size = Pt(12)
    run.bold = True

    # ── 왜 만들었나? ──
    style_heading(doc, "1. 왜 만들었나?", 1)
    doc.add_paragraph(
        "• 연구원들이 논문을 찾고 분석하는 데 많은 시간을 소비\n"
        "• 수천 편의 논문에서 원하는 정보를 빠르게 찾기 어려움\n"
        "• 논문 간 트렌드, 비교, 아이디어 발굴은 사람이 하기에 한계가 있음\n\n"
        "→ AI가 24시간 대기하면서 연구원의 질문에 즉시 답변해주는 '논문 전문 비서'를 만들었습니다."
    )

    # ── 무엇을 할 수 있나? ──
    style_heading(doc, "2. 무엇을 할 수 있나?", 1)
    doc.add_paragraph("이 시스템은 14가지 종류의 질문에 답할 수 있습니다:")
    doc.add_paragraph()

    add_table(doc,
        ["분류", "할 수 있는 일", "질문 예시"],
        [
            ["논문 검색", "원하는 논문을 찾아 내용 요약", "\"OLED 수명 관련 최신 논문 알려줘\""],
            ["문헌 리뷰", "특정 주제의 연구 동향 정리", "\"마이크로LED 전사 기술 리뷰해줘\""],
            ["심층 분석", "특정 논문의 핵심 내용 분석", "\"이 논문의 실험 방법 설명해줘\""],
            ["통계 분석", "논문 수, 연도별 추이 등 숫자 분석", "\"2024년 OLED 논문 몇 편이야?\""],
            ["아이디어 제안", "논문들을 조합해 새 연구 방향 제시", "\"양자점과 OLED 결합 아이디어 있어?\""],
            ["트렌드 분석", "기술 분야별 변화 추이 분석", "\"최근 3년 디스플레이 기술 트렌드는?\""],
            ["실험 설계", "문헌 기반 실험 계획 제안", "\"페로브스카이트 LED 효율 실험 설계해줘\""],
            ["경쟁사 분석", "경쟁사 연구 동향 파악", "\"삼성 vs LG 디스플레이 연구 비교\""],
            ["보고서 작성", "연구 보고서 초안 자동 생성", "\"OLED 수명 개선 보고서 초안 만들어줘\""],
        ],
        col_widths=[2.5, 5, 6],
    )

    # ── 어떻게 동작하나? ──
    style_heading(doc, "3. 어떻게 동작하나? (간단 설명)", 1)
    doc.add_paragraph(
        "① 연구원이 채팅창에 질문을 입력합니다\n\n"
        "② AI가 질문을 분석합니다\n"
        "   - 날짜 관련 표현이 있으면 기간을 자동 계산\n"
        "     (예: \"최근 6개월\" → 2025년 10월 ~ 2026년 3월)\n"
        "   - 질문 유형을 파악 (검색? 통계? 아이디어? 등)\n\n"
        "③ 14명의 'AI 전문가' 중 적합한 1명이 배정됩니다\n\n"
        "④ 배정된 AI 전문가가 논문 DB를 검색하고 답변을 생성합니다\n\n"
        "⑤ 출처 논문 목록과 함께 답변이 전달됩니다"
    )

    # 다이어그램 삽입
    if os.path.exists(DIAGRAM_PATH):
        doc.add_paragraph()
        style_heading(doc, "시스템 구성도", 2)
        doc.add_picture(DIAGRAM_PATH, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 현재 성과 ──
    style_heading(doc, "4. 현재 성과", 1)
    add_table(doc,
        ["항목", "결과", "의미"],
        [
            ["논문 데이터", "1,614편 적재", "SID/Wiley 디스플레이 논문 DB 구축 완료"],
            ["날짜 이해력", "100% 정확", "\"작년 여름\", \"최근 3개월\" 등 자연어 날짜를 완벽 이해"],
            ["질문 분류", "50% (경량 AI 기준)", "고성능 AI 전환 시 대폭 향상 예상"],
            ["전체 답변 품질", "70% (샘플 테스트)", "기본 질문에는 충분히 활용 가능한 수준"],
        ],
        col_widths=[3, 3.5, 7],
    )

    # ── 사용 환경 ──
    style_heading(doc, "5. 사용 환경 및 보안", 1)
    doc.add_paragraph(
        "• 웹 브라우저(Open WebUI)에서 ChatGPT처럼 대화형으로 사용\n"
        "• 사내 서버에서 독립 운영 — 외부 클라우드 불필요\n"
        "• 폐쇄망(인터넷 차단 환경)에서도 설정 변경만으로 바로 사용 가능\n"
        "• 모든 데이터가 사내에 머물러 정보 유출 우려 없음"
    )

    # ── 기대 효과 ──
    style_heading(doc, "6. 기대 효과", 1)
    add_table(doc,
        ["영역", "기존", "AI 도입 후"],
        [
            ["논문 검색", "직접 DB 검색 + 수작업 필터링\n(30분~수시간)", "자연어 질문으로 즉시 검색\n(수 초)"],
            ["문헌 리뷰", "수십 편 논문 직접 읽고 정리\n(수일)", "AI가 핵심 내용 자동 요약\n(수 분)"],
            ["트렌드 파악", "수동 집계 + 차트 작성\n(수시간)", "\"최근 3년 트렌드\" 질문 한 마디\n(수 초)"],
            ["아이디어 발굴", "개인 경험에 의존", "1,600편 교차 분석 기반 제안"],
            ["보고서 작성", "백지에서 시작", "AI가 초안 생성 → 연구원이 수정"],
        ],
        col_widths=[2.5, 5.5, 5.5],
    )

    # ── 향후 계획 ──
    style_heading(doc, "7. 향후 계획", 1)
    add_table(doc,
        ["단계", "내용", "효과"],
        [
            ["1단계\n(현재)", "경량 AI로 기본 시스템 구축 완료", "기본 질문 응답 가능"],
            ["2단계\n(단기)", "고성능 AI(235B 모델)로 전환", "질문 분류 정확도 50%→90%+ 향상"],
            ["3단계\n(중기)", "폐쇄망 배포 + 사용자 교육", "전 연구원 실무 활용 시작"],
            ["4단계\n(장기)", "논문 자동 업데이트 + 에이전트 확장", "최신 논문 실시간 반영"],
        ],
        col_widths=[2, 6, 5.5],
    )

    # ── 요약 ──
    style_heading(doc, "요약", 1)
    doc.add_paragraph(
        "Co-Scientist Agent는 연구원의 논문 검색·분석 업무를 AI로 자동화하는 시스템입니다. "
        "현재 1,614편의 디스플레이 논문을 기반으로 14가지 유형의 질문에 답변할 수 있으며, "
        "사내 서버에서 독립 운영되어 정보 보안 우려가 없습니다.\n\n"
        "고성능 AI 모델 전환과 폐쇄망 배포를 통해, "
        "전 연구원이 일상적으로 활용할 수 있는 'AI 논문 비서'로 발전시킬 계획입니다."
    )

    path = os.path.join(DOCS_DIR, "경영진_보고서_AI_논문_분석_비서.docx")
    doc.save(path)
    print(f"Executive report saved: {path}")


if __name__ == "__main__":
    generate_technical_report()
    generate_executive_report()
